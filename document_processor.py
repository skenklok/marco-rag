import os
import logging
from PyPDF2 import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer
import numpy as np
import faiss
import torch


# Initialize logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Initialize the sentence transformer model for embeddings
model = SentenceTransformer('all-MiniLM-L6-v2')
logging.info("Loaded Sentence Transformer model for embeddings.")

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    text = ''
    try:
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            text += page.extract_text() + '\n'
    except Exception as e:
        logging.error(f"Error reading {pdf_path}: {e}")
    return text

def extract_text_from_docx(docx_path):
    """Extract text from a Word document."""
    text = ''
    try:
        doc = Document(docx_path)
        text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    except Exception as e:
        logging.error(f"Error reading {docx_path}: {e}")
    return text

def get_embeddings(texts, model):  # Add 'model' as an argument
    """
    Generate embeddings for a list of text documents.
    
    Parameters:
    - texts: List of text documents.
    - model: SentenceTransformer model used for generating embeddings.
    
    Returns:
    - embeddings_array: Array containing embeddings for the input texts.
    """
    embeddings = []
    for text in texts:
        embedding = model.encode([text])[0]
        embeddings.append(embedding)
    embeddings_array = np.array(embeddings)
    logging.info(f"Shape of embeddings array: {embeddings_array.shape}")  # Add console log for shape
    return embeddings_array


def process_documents(folder_path, model):
    texts = []
    document_embeddings = []

    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path) and file_path.endswith('.docx'):
            document = Document(file_path)
            text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
            texts.append(text)

            with torch.no_grad():
                doc_embeddings = model.model.get_input_embeddings()(
                    model.tokenizer(text, return_tensors="pt")["input_ids"].detach().clone()
                ).cpu().numpy()
                aggregated_embedding = np.mean(doc_embeddings, axis=1).squeeze(0)
                document_embeddings.append(aggregated_embedding)

    return texts, np.array(document_embeddings)



def create_faiss_index(embeddings):
    # Check if there are any embeddings
    if len(embeddings) == 0:
        raise ValueError("No embeddings found")

    embeddings_array = np.array(embeddings)

    # Check if there are any numerical embeddings
    if embeddings_array.size == 0:
        raise ValueError("No numerical embeddings found")

    # Ensure embeddings array is 2-dimensional
    if len(embeddings_array.shape) == 1:
        embeddings_array = embeddings_array.reshape(1, -1)

    # Create a Faiss index
    d = embeddings_array.shape[1]  # Dimensionality of the embeddings
    index = faiss.IndexFlatL2(d)

    # Add the embeddings to the index
    index.add(embeddings_array)

    return index

